import re
import requests
import time
from fake_useragent import UserAgent
import docx
from docx.shared import Inches
from lxml import etree

startTime = time.time()


# 下载图片保存
def getPicture(url, headers, proxy):
    with open('Temp.png', 'wb+')as f:
        f.write(requests.get(url, headers=headers, proxies=proxy).content)


url = 'https://mp.weixin.qq.com/s?__biz=MzU2NzEwMDc1MA==&mid=2247493453&idx=1&sn=c9dc7ba1cf0cdc49f14d380e299afbf9&chksm=fca0f36ccbd77a7aac6aea831eddf7556de038b60a2197f3a1a6a5b1dff563c1234f6c189764#rd'
HttpIp = '58.218.214.136:4638'
proxy = {'http': HttpIp, 'https': HttpIp}
headers = {
    'User-Agent': UserAgent().random,
}

doc_new = docx.Document()

selector = etree.HTML(requests.get(url, headers=headers, proxies=proxy).text)

contentAll = selector.xpath('//div[@id="js_content"]/p')[2:][0:-4]
print(bool(not contentAll))
for i in selector.xpath('//div[@id="js_content"]/p')[2:][0:-4]:
    text = etree.tostring(i, encoding='utf-8').decode()
    print(text)
    if 'data-src' in text:  # 图片
        tempURL = re.findall('data-src="(https.*?)"', text)
        if tempURL:
            pictureUrl = tempURL[0]
            # 下载图片
            getPicture(pictureUrl, headers, proxy)
            # 添加图片
            doc_new.add_picture('Temp.png', width=Inches(6))
    else:  # 文字内容
        contentList = etree.HTML(text).xpath('//p/text()')
        if contentList:
            content = contentList[0]
            # 添加到docx中
            doc_new.add_paragraph(content)
        else:  # 标题
            if '<p><br/></p>' != text:
                if 'max-width: 100%;color: rgb(255, 255, 255);' in text:
                    title1 = ''.join(re.findall('[\u4e00-\u9fa5]+', text))
                    # 添加到docx中
                    doc_new.add_heading(title1, level=1)
                else:
                    title2 = ''.join(re.findall('[\u4e00-\u9fa5]+', text))
                    if title2:  # 不为空才添加
                        # 添加到docx中
                        doc_new.add_paragraph(title2, style='Intense Quote')

            else:
                doc_new.add_paragraph()

# 保存文档
doc_new.save('微信公众号/article/1.docx')
print('耗时：', time.time() - startTime)
