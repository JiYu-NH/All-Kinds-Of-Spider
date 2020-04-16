import time, re
import random
import datetime
from queue import Queue
import requests
from lxml import etree
from fake_useragent import UserAgent
import docx
from docx.shared import Inches
import threading

# 抓取微信公众号所有文章链接
"""
for begin in range(355, 970, 5):
    print('\n\n' + str(begin) + '\n\n')
    # 微信公众号平台 -- 素材 -- 超链接 -- 公众号文章   抓取出来的
    # begin/5+1 为 页数，第一页的begin为0：
    url = "https://mp.weixin.qq.com/cgi-bin/appmsg?action=list_ex&begin=%s&count=5&fakeid=MzU2NzEwMDc1MA==&type=9&query=&token=578630500&lang=zh_CN&f=json&ajax=1" % begin
    HttpIp = '58.218.214.156:8905'
    proxy = {'http': HttpIp, 'https': HttpIp}
    headers = {
        'User-Agent': "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_9_3) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/35.0.1916.47 Safari/537.36",
        # 抓 url 时，直接复制过来使用
        'Cookie': 'ua_id=g55dHiQIwbBe3alQAAA...'
    }

    # 保存抓到的文章链接
    with open('urls.txt', 'a+', encoding='utf-8') as f:
        for aid in requests.get(url, headers=headers, proxies=proxy).json()['app_msg_list']:
            # 标题、日期、文章链接
            data = '%s@%s@%s' % (
                aid['title'], datetime.datetime.fromtimestamp(aid['update_time']).strftime('%Y-%m-%d'), aid['link'])
            print(data)
            f.write(data + '\n')
    # 间隔25-40秒、太快要被封
    time.sleep(random.randint(30, 70))
"""


# 提取IP
def ExtractIP(ipQueue: Queue):
    res = requests.get('.....').json()
    if res['success']:
        print('代理余额：', res['data'])
    while 1:
        print('提取IP中...')
        url = '...'
        res = requests.get(url).json()
        # 不成功则打印信息
        if res['success'] == 'false':
            print(res['msg'])
            # 添加白名单
            if '白名单IP' in res['msg']:
                myIP = res['msg'].split('登录')[0]
                addIpUrl = '...' + myIP
                time.sleep(2)
                print(requests.get(addIpUrl).json()['msg'])
        else:
            datas = res['data']
            for data in datas:
                ipQueue.put(str(data['IP']) + ":" + str(data['Port']))
        # 隔两秒再次提取
        time.sleep(2)
        if ipQueue.qsize() > 150: break


# 读取链接
def printURL(taskQueue: Queue):
    with open('urls.txt', encoding='utf-8')as f:
        for i in f.read().split('\n'):
            taskQueue.put(i)


def Grab(taskQueue: Queue, ipQueue: Queue):
    # 队列里有任务时
    while not taskQueue.empty():
        jinruData = 1
        taskNum = taskQueue.qsize()
        urlData = taskQueue.get()
        # for urlData, num in zip(urlList, range(1, lenURL + 1)):
        if urlData:
            startTime = time.time()
            titleURL, dateURL, url = urlData.split('@')
            # print('====================', titleURL, '资讯汇总' not in titleURL)
            if '资讯汇总' not in titleURL:
                # 新建文档
                doc_new = docx.Document()
                # doc_new = docQueue.get()
                while 1:
                    try:
                        jinruData = 2
                        HttpIp = ipQueue.get()
                        proxy = {'http': HttpIp, 'https': HttpIp}
                        headers = {
                            'User-Agent': UserAgent().random,
                        }
                        selector = etree.HTML(requests.get(url, headers=headers, proxies=proxy, timeout=3).text)
                        # print('本次IP：', HttpIp)
                        contentAll = selector.xpath('//div[@id="js_content"]/p')[2:][0:-4]

                        jinruData = 3
                        # 文章被封，没有数据，继续下一个队列中的任务
                        if not contentAll: break
                        #     print('没有抓到数据没有抓到数据没有抓到数据%s' % contentAll)
                        #     break
                        # else:
                        #     print('抓到' * 30)
                        jinruData = 4
                        for i in contentAll:
                            text = etree.tostring(i, encoding='utf-8').decode()
                            jinruData = 5
                            if 'data-src' in text:  # 图片
                                tempUrl = re.findall('data-src="(https://mmbiz.qpic.cn.*?)"', text)
                                if tempUrl:
                                    pictureUrl = tempUrl[0]
                                    # print(pictureUrl)
                                    # 下载图片
                                    fileName = 'tempPicture\p' + str(int(time.time() * 100000)) + str(
                                        random.randint(1, 9999)) + '.png'
                                    if (getPicture(pictureUrl, headers, proxy, fileName)):
                                        # 添加图片
                                        doc_new.add_picture(fileName, width=Inches(6))
                            else:  # 文字内容
                                contentList = etree.HTML(text).xpath('//p/text()')
                                if contentList:
                                    content = contentList[0]
                                    # 添加到docx中
                                    doc_new.add_paragraph(content)
                                else:  # 标题
                                    if '<p><br/></p>' != text:
                                        if 'max-width: 100%;color: rgb(255, 255, 255);' in text:
                                            title1 = ''.join(re.findall('[\u4e00-\u9fa5]+', text))
                                            # 添加到docx中
                                            doc_new.add_heading(title1, level=1)
                                        else:
                                            title2 = ''.join(re.findall('[\u4e00-\u9fa5]+', text))
                                            if title2:  # 不为空才添加
                                                # 添加到docx中
                                                doc_new.add_paragraph(title2, style='Intense Quote')
                                    else:
                                        doc_new.add_paragraph()
                        # 保存文档
                        doc_new.save('article/%s_%s-%s.docx' % (dateURL, titleURL, random.randint(1, 999)))
                        # print('article/%s_%s.docx' % (dateURL, titleURL), '已保存' * 10)
                        # 正常退出
                        break
                    except Exception as e:
                        print('*' * 20, '异常信息：', e)
                        HttpIp = ipQueue.get()
            print('本次耗时：%s\n标题：%s\n资讯汇总：%s\n进度：%s\nurl:%s\n 剩余任务数量：%s\n剩余IP数量：%s\n%s_%s.docx%s' % (
                time.time() - startTime, titleURL, bool('资讯汇总' not in titleURL), jinruData, url, taskNum,
                ipQueue.qsize(), dateURL, titleURL, "\n\n"))

    print('没有任务了' * 10)
    return


# 下载图片保存
def getPicture(url, headers, proxy, fileName):
    with open(fileName, 'wb+')as f:
        data = requests.get(url, headers=headers, proxies=proxy).content
        if b'<html>' in data or b'<head>' in data:
            return False
        else:
            f.write(data)
            return True


# 任务队列
taskQueue = Queue()
# IP队列
ipQueue = Queue()
# doc队列
# docQueue = Queue()

# for _ in range(1000):
#     print('添加doc', _)
#     docQueue.put(docx.Document())

ExtractIP(ipQueue)

printURL(taskQueue)

t_list = []

for i in range(20):
    t = threading.Thread(target=Grab, args=(taskQueue, ipQueue))
    t_list.append(t)

for t in t_list:
    t.start()
for t in t_list:
    t.join()

print('全部结束：', taskQueue.qsize())
